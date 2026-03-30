#!/usr/bin/env python3
"""
CV DOCX Generator — renders CV from JSON data following the fixed template structure.
Usage: python3 make_cv_docx.py <input.json> <output.docx>
"""
import sys, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def make_cv(data, output_path):
    doc = Document()

    # Page margins — tight to fit 2 pages (matching original CV style)
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Remove default paragraph spacing
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    def set_spacing(p, before=0, after=2):
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)

    # ── Section heading: bold + underline + keep-with-next ─
    def add_heading(text):
        p = doc.add_paragraph()
        set_spacing(p, before=6, after=2)
        # Keep heading with next paragraph — never orphan heading at page bottom
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.underline = True
        run.font.size = Pt(11)

    # ── Bullet with optional bold label ────────────────────
    def add_bullet(label=None, text=''):
        p = doc.add_paragraph(style='List Bullet')
        set_spacing(p, before=0, after=0)
        if label:
            r1 = p.add_run(label + ': ')
            r1.bold = True
            r1.font.size = Pt(11)
            r2 = p.add_run(text)
            r2.font.size = Pt(11)
        else:
            r = p.add_run(text)
            r.font.size = Pt(11)

    # ══ 1. HEADER ══════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=2)
    r = p.add_run(data['name'])
    r.bold = True
    r.font.size = Pt(14)

    for line in data['contacts']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, after=1)
        r = p.add_run(line)
        r.font.size = Pt(11)

    # LinkedIn as real clickable hyperlink
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=6)
    url = data['linkedin']
    # Add relationship
    part = p.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    # Build hyperlink XML element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')
    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Blue color like standard Word hyperlink
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    # Font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '22')  # 11pt = 22 half-points
    rPr.append(sz)
    run_elem.append(rPr)
    t = OxmlElement('w:t')
    t.text = url
    run_elem.append(t)
    hyperlink.append(run_elem)
    p._p.append(hyperlink)

    # ══ 2. PERSONAL PROFILE ════════════════════════════════
    add_heading('PERSONAL PROFILE')
    p = doc.add_paragraph()
    set_spacing(p, after=6)
    r = p.add_run(data['personal_profile'])
    r.font.size = Pt(11)

    # ══ 3. KEY SKILLS ══════════════════════════════════════
    add_heading('KEY SKILLS')
    for skill in data['key_skills']:
        add_bullet(label=skill['label'], text=skill['text'])

    # ══ 4. WORK EXPERIENCE ═════════════════════════════════
    add_heading('WORK EXPERIENCE')
    for job in data['work_experience']:
        # Date + Company on one line, bold
        p = doc.add_paragraph()
        set_spacing(p, before=4, after=0)
        r = p.add_run(f"{job['date']}   {job['company']}")
        r.bold = True
        r.font.size = Pt(11)
        # Job title, bold, indented
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.8)
        set_spacing(p2, before=0, after=2)
        r2 = p2.add_run(job['title'])
        r2.bold = True
        r2.font.size = Pt(11)
        # Bullets
        for b in job['bullets']:
            add_bullet(text=b)

    # ══ 5. EDUCATION AND TRAINING ══════════════════════════
    add_heading('EDUCATION AND TRAINING')
    # Invisible 2-column table
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    # Remove all borders
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    for edu in data['education']:
        row = table.add_row()
        # Left cell: date bold
        c0 = row.cells[0]
        c0.width = Cm(3.5)
        p0 = c0.paragraphs[0]
        set_spacing(p0, before=2, after=2)
        r0 = p0.add_run(edu['date'])
        r0.bold = True
        r0.font.size = Pt(11)
        # Right cell: institution normal + qualification bold
        c1 = row.cells[1]
        p1 = c1.paragraphs[0]
        set_spacing(p1, before=2, after=0)
        r1 = p1.add_run(edu['institution'])
        r1.font.size = Pt(11)
        p2 = c1.add_paragraph()
        set_spacing(p2, before=0, after=2)
        r2 = p2.add_run(edu['qualification'])
        r2.bold = True
        r2.font.size = Pt(11)

    doc.add_paragraph()

    # ══ 6. ADDITIONAL INFORMATION ══════════════════════════
    add_heading('ADDITIONAL INFORMATION')
    for item in data['additional_info']:
        add_bullet(label=item['label'], text=item['text'])

    doc.save(output_path)
    print(f'Saved: {output_path}')
    # Estimate page count via paragraph count (rough check)
    paras = len(doc.paragraphs)
    print(f'Paragraphs: {paras} (rough estimate — open in Word to verify page count)')


# ── CV DATA FOR TRICEL SOFTWARE ENGINEER ───────────────
cv_data = {
    "name": "Serhii Baliasnyi",
    "contacts": [
        "Ballydowney, Killarney, Co Kerry",
        "Mobile: 0852007612",
        "E-mail: sergey070373@gmail.com"
    ],
    "linkedin": "https://www.linkedin.com/in/serhii-baliasnyi-290b72246/",
    "personal_profile": (
        "Versatile Software Engineer and IT professional with 29+ years of experience designing, "
        "developing and maintaining robust software solutions across enterprise environments. "
        "Proficient in Java, Python, Docker, Linux and Azure with hands-on CI/CD pipeline experience. "
        "Currently developing full-stack AI-powered platforms combining FastAPI backends with "
        "React/TypeScript frontends. Based in Killarney, I am immediately available for the "
        "Software Engineer role at Tricel Head Office and eager to contribute to your growing team."
    ),
    "key_skills": [
        {"label": "Languages & Frameworks", "text": "Java, Spring Boot, Python (FastAPI, Django), JavaScript, TypeScript, React, Node.js, PHP (Symfony), HTML5, CSS."},
        {"label": "Databases", "text": "PL/SQL, PostgreSQL, MySQL, MS SQL Server, MongoDB, Firebird, Supabase."},
        {"label": "Infrastructure & DevOps", "text": "Docker, Linux (Ubuntu, CentOS), Microsoft Azure, AWS, CI/CD pipelines, Git, VMware, VirtualBox."},
        {"label": "Cloud & AI", "text": "RAG systems, LLM integrations (OpenAI, Claude), LlamaIndex, Vector Databases, n8n Workflows."},
        {"label": "Support & Operations", "text": "System administration, troubleshooting, GDPR compliance, network configuration, hardware maintenance."},
        {"label": "Tools", "text": "SharePoint Online, SPFx, Microsoft Graph API, Power Apps, Power Automate, Jira, Agile concepts."},
    ],
    "work_experience": [
        {
            "date": "Aug 2025 - Present",
            "company": "Independent Tech Consultant & Developer",
            "title": "Full-Stack Software Engineer / AI Developer",
            "bullets": [
                "Sole architect and developer of AI MediaFlow (aimediaflow.net) — full-stack AI content automation platform.",
                "Built React/TypeScript frontend, FastAPI backend, Supabase (PostgreSQL + Vector storage), deployed via Docker and Vercel.",
                "Designed RAG systems with LLM integrations (OpenAI, Claude); implemented document workflow automation.",
                "End-to-end CI/CD pipeline management on Hetzner VPS with Docker containers.",
            ]
        },
        {
            "date": "Jun 2024 - Present",
            "company": "KPFA (part time)",
            "title": "System Computer Administrator",
            "bullets": [
                "Developed custom SharePoint SPFx web parts in TypeScript/React: Staff Record System with Excel export, modular architecture (8 handler modules).",
                "Designed and implemented RAG document indexing system using FastAPI and LlamaIndex.",
                "Process automation and electronic document workflow support using Power Apps and Power Automate.",
                "Technical support, hardware maintenance, user access management and GDPR compliance.",
            ]
        },
        {
            "date": "Aug 2003 - Mar 2022",
            "company": "Ukraine Government / State Enterprise Information Court System",
            "title": "Chief of Technical Support Department, Kherson Region",
            "bullets": [
                "Directed daily operations of regional IT department for 19 years; managed staff and resources.",
                "Deployed and maintained critical infrastructure: servers, networks, workstations across government locations.",
                "Administered MSSQL, MySQL, Firebird, PostgreSQL databases; configured Cisco, VPN, software firewalls.",
                "Managed Git, Docker, VMware/VirtualBox environments; ensured GDPR and data security compliance.",
                "Developed internal tools using HTML, JavaScript, ReactJS, NodeJS.",
            ]
        },
        {
            "date": "Sept 1999 - Aug 2003",
            "company": "Ukraine Government / State Enterprise Information Court System",
            "title": "Engineer of Technical Support Department",
            "bullets": [
                "Installed and maintained IT systems; administered MSSQL, MySQL, Firebird and PostgreSQL servers.",
                "Troubleshooting, strategic testing, website support and GDPR compliance.",
            ]
        },
        {
            "date": "Mar 1998 - Sept 1999",
            "company": "DnieperStyle – Ukraine Private Computer Company",
            "title": "Technical Support Engineer, PHP Programmer",
            "bullets": [
                "Website development: HTML, JavaScript, PHP, Perl, MySQL, PostgreSQL; remote server administration.",
            ]
        },
    ],
    "education": [
        {
            "date": "Oct 2022 - Present",
            "institution": "eCollege, Dublin",
            "qualification": "Java Foundation (1z0-811), Professional Java Associate (1z0-808)"
        },
        {
            "date": "Feb 2019 - July 2019",
            "institution": "Online Courses, University of Michigan",
            "qualification": "Introduction to SQL, Building Database Applications in PHP, Building Web Applications in PHP"
        },
        {
            "date": "Sep 2009 - May 2011",
            "institution": "Kherson State University, Ukraine",
            "qualification": "Bachelor of Economics"
        },
        {
            "date": "Sep 1990 - June 1995",
            "institution": "Sevastopol State University, Ukraine",
            "qualification": "Bachelor of Microelectronics"
        },
    ],
    "additional_info": [
        {"label": "Local Candidate", "text": "Residing in Killarney, available for immediate start and onsite work."},
        {"label": "Key Projects", "text": "Founder of AI MediaFlow (aimediaflow.net) — AI content automation platform; developer of EnVocab (Android vocabulary app on Google Play)."},
        {"label": "Automation Focus", "text": "Experienced in building custom workflows, automation tools and CI/CD pipelines using Python, Docker and cloud platforms."},
        {"label": "Languages", "text": "English (B2), Ukrainian, Russian."},
    ]
}

if __name__ == '__main__':
    out = sys.argv[1] if len(sys.argv) > 1 else '/opt/yt-api/cv/cv_tricel_software_engineer.docx'
    make_cv(cv_data, out)
